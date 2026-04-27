import os
import pandas as pd
from docx import Document

print("--- EXPLORING 2022_SCHEME EXCEL ---")
scheme_path = '2022_SCHEME_3 SEM_22AI34_Data Structures with Applications_Calculation_Reshma S-Print (1).xlsx'
try:
    xl = pd.ExcelFile(scheme_path)
    print("Sheets:", xl.sheet_names)
    if 'Theory' in xl.sheet_names:
        df = xl.parse('Theory', nrows=20, header=None) # parse without header to see raw layout
        print("First 20 rows of Theory sheet:")
        for idx, row in df.iterrows():
            print(f"Row {idx}:", [str(x)[:20] if not pd.isna(x) else '' for x in row.values])
except Exception as e:
    print("Error reading scheme:", e)

print("\n--- EXPLORING DOCX ---")
docx_path = 'AAI-CIA1-26-08-2025-set 1.docx'
try:
    doc = Document(docx_path)
    print("Paragraphs in docx:")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"P{i}: {text[:100]}")
    print("\nTables in docx:")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for j, row in enumerate(table.rows):
            print(f"  Row {j}: {[cell.text.strip() for cell in row.cells]}")
except Exception as e:
    print("Error reading docx:", e)

print("\n--- EXPLORING IA_MARKS XLS ---")
xls_path = 'IA_Marks (4).xls'
try:
    try:
        # It might be HTML disguised as xls, let's try reading as HTML first
        dfs = pd.read_html(xls_path)
        print("XLS parsed as HTML table. Found", len(dfs), "tables.")
        for i, df in enumerate(dfs):
            print(f"Table {i} head:")
            print(df.head())
    except ValueError:
        # Try as actual excel
        xl = pd.ExcelFile(xls_path, engine='xlrd')
        print("Sheets in xls:", xl.sheet_names)
        df = xl.parse(xl.sheet_names[0], nrows=10, header=None)
        print("First 10 rows:")
        for idx, row in df.iterrows():
            print(f"Row {idx}:", [str(x)[:20] if not pd.isna(x) else '' for x in row.values])
except Exception as e:
    print("Error reading xls:", e)
