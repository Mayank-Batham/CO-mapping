from docx import Document

doc_path = 'AAI-CIA1-26-08-2025-set 1.docx'
try:
    doc = Document(doc_path)
    print("Found", len(doc.tables), "tables in QP.")
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx}:")
        for r_idx, row in enumerate(table.rows):
            print(f"  Row {r_idx}: {[cell.text.strip().replace('\n', ' ') for cell in row.cells]}")
except Exception as e:
    print("Error reading docx:", e)
